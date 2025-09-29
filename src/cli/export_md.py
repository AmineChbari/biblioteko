import os
import sys
import base64
import json
import requests
import fitz  # PyMuPDF library
from dotenv import load_dotenv
# --- NOUVELLE DÉPENDANCE REQUISE ---
# Pour créer un fichier .docx, nous devons utiliser la librairie 'python-docx'.
# Vous devez l'installer : poetry add python-docx
from docx import Document 

# --- NOUVEAU PROMPT AMÉLIORÉ ---
PROMPT_OCR = (
    "Extract all text from this scanned document page. "
    "The text extracted must be in docx format with titles, subtitles and backlines and all the display elements."
    "Respond only with the raw docx content, without any additional commentary or explanation."
)
MODEL_API_URL = "https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-flash-preview-05-20:generateContent"
# ------------------------------

def get_text_from_image(image_data, api_key):
    """
    Sends image data to the Gemini API for text extraction (OCR).

    Args:
        image_data (bytes): The raw image data.
        api_key (str): Your Gemini API key.

    Returns:
        str: The extracted text or an error message.
    """
    # The API URL of the model
    api_url = f"{MODEL_API_URL}?key={api_key}"

    # Convert the image data to a Base64 string
    encoded_string = base64.b64encode(image_data).decode("utf-8")

    # Construct the API payload
    payload = {
        "contents": [
            {
                "parts": [
                    {
                        "text": PROMPT_OCR
                    },
                    {
                        "inlineData": {
                            "mimeType": "image/jpeg",
                            "data": encoded_string
                        }
                    }
                ]
            }
        ]
    }

    headers = {
        "Content-Type": "application/json"
    }

    try:
        response = requests.post(api_url, headers=headers, data=json.dumps(payload))
        response.raise_for_status()  # Raise an exception for bad status codes (4xx or 5xx)
        response_data = response.json()

        # Check for candidates and text
        if "candidates" in response_data and response_data["candidates"]:
            candidate = response_data["candidates"][0]
            if "content" in candidate and "parts" in candidate["content"]:
                for part in candidate["content"]["parts"]:
                    if "text" in part:
                        return part["text"]
        return "No text extracted."

    except requests.exceptions.HTTPError as err:
        return f"HTTP error occurred: {err}"
    except Exception as err:
        return f"An error occurred: {err}"

def write_to_docx(filename, pages_content):
    """
    Writes extracted text pages to a DOCX file using python-docx.
    Each page is treated as a new section or set of paragraphs.
    
    NOTE: Since Gemini retourne du Markdown, cette fonction ajoute le texte brut 
    (contenant la syntaxe Markdown) sous forme de paragraphes simples dans le DOCX.
    """
    document = Document()
    
    document.add_heading('Texte Extrait des Scans de Livre', 0)

    for i, page_text in enumerate(pages_content):
        # Ajoute le titre de la page
        document.add_heading(f'Page {i + 1}', level=1)
        
        # Ajoute le contenu. Nous coupons les doubles sauts de ligne pour simuler des paragraphes
        paragraphs = page_text.split('\n\n')
        
        for paragraph in paragraphs:
            if paragraph.strip():
                # Ajoute le texte brut du paragraphe
                document.add_paragraph(paragraph.strip())
        
        # Ajoute un séparateur visuel entre les pages
        document.add_page_break()

    document.save(filename)


def main():
    """
    Main function to process all pages of a PDF file.
    """
    # --- Configuration ---
    # The name of the output DOCX file. CHANGED from .md to .docx
    output_file = "book_text.docx"
    
    # Path to the .env file.
    script_dir = os.path.dirname(os.path.abspath(__file__))
    env_file_path = os.path.join(script_dir, '../../.env')

    # Load environment variables from .env file
    load_dotenv(dotenv_path=env_file_path)

    # Get the API key from the environment
    api_key = os.getenv("GEMINI_API_KEY")
    if not api_key:
        print("Error: GEMINI_API_KEY not found. Please ensure it is set in your .env file.")
        print(f"Looked for .env file at: {env_file_path}")
        return

    # Get the PDF file path from the command-line arguments
    if len(sys.argv) < 2:
        print("Usage: python scan_to_markdown.py <path_to_pdf_file>")
        return

    input_pdf_file = sys.argv[1]

    # Check if the PDF file exists
    if not os.path.exists(input_pdf_file):
        print(f"Error: PDF file not found at '{input_pdf_file}'.")
        return

    extracted_pages = []
    
    try:
        pdf_document = fitz.open(input_pdf_file)
        num_pages = pdf_document.page_count
        
        print(f"Found {num_pages} page(s) in the PDF file.")

        # Limitez le traitement à quelques premières pages pour les tests rapides
        pages_to_process = min(num_pages, 5) 
        
        for i in range(pages_to_process):
            print(f"Processing page {i+1}/{num_pages}...")
            
            # Render the page to a high-resolution image in memory
            page = pdf_document.load_page(i)
            pix = page.get_pixmap(matrix=fitz.Matrix(300 / 72, 300 / 72))  # Render at 300 DPI
            image_data = pix.tobytes("jpeg")
            
            text = get_text_from_image(image_data, api_key)
            
            if text.startswith(("HTTP error", "Error encoding")):
                print(f"Failed to process page {i+1}: {text}")
                # Nous arrêtons ici pour le débogage de la première page
                return
            else:
                print(f"Successfully extracted text from page {i+1}.")
                extracted_pages.append(text)

        pdf_document.close()

    except Exception as e:
        print(f"An error occurred while processing the PDF: {e}")
        return

    # Write all the extracted text to a single DOCX file
    try:
        print(f"\nWriting {len(extracted_pages)} pages to {output_file}...")
        write_to_docx(output_file, extracted_pages)
    except NameError:
        print("\nERREUR: La librairie 'python-docx' n'est pas installée.")
        print("Veuillez installer cette dépendance avant de lancer le script, en utilisant :")
        print("poetry add python-docx")
        return
    except Exception as e:
        print(f"\nUne erreur inattendue s'est produite lors de la création du fichier DOCX : {e}")
        return

    print("\nProcessing complete!")
    print(f"All extracted text has been saved to '{output_file}'.")

if __name__ == "__main__":
    main()
